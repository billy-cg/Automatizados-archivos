import tkinter as tk
from tkinter import filedialog, messagebox
from tkinterdnd2 import TkinterDnD, DND_FILES

import pandas as pd
from pathlib import Path
import os
import threading

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from docx import Document
from PyPDF2 import PdfReader


# =============================
# TEMAS
# =============================

TEMA_OSCURO = {
    "BG": "#1e1e1e",
    "FG": "#d4d4d4",
    "BTN_BG": "#2d2d2d",
    "BTN_FG": "#ffffff",
    "ENTRY_BG": "#252526",
    "LIST_BG": "#252526",
    "ACCENT": "#007acc",
    "DROP_BG": "#252526",
    "DROP_FG": "#9cdcfe"
}

TEMA_CLARO = {
    "BG": "#f2f2f2",
    "FG": "#000000",
    "BTN_BG": "#e0e0e0",
    "BTN_FG": "#000000",
    "ENTRY_BG": "#ffffff",
    "LIST_BG": "#ffffff",
    "ACCENT": "#007acc",
    "DROP_BG": "#ffffff",
    "DROP_FG": "#333333"
}


# =============================
# APP
# =============================

class AutomatizadorApp:

    def __init__(self, root):
        self.root = root
        self.root.title("Automatizador de Datos")
        self.root.geometry("600x700")

        self.tema = TEMA_OSCURO
        self.archivo_seleccionado = None
        self.ultimo_archivo = None

        # -----------------------------
        # BOTÓN TEMA (ESQUINA SUP IZQ)
        # -----------------------------
        self.btn_tema = tk.Button(
            root,
            text="Cambiar tema",
            command=self.toggle_tema
        )
        self.btn_tema.place(x=10, y=10)

        # -----------------------------
        # ARCHIVO ENTRADA
        # -----------------------------
        self.lbl_titulo = tk.Label(root, text="Archivo de entrada", font=("Arial", 13))
        self.lbl_titulo.pack(pady=40)

        self.frame_archivo = tk.Frame(root)
        self.frame_archivo.pack(pady=4)

        self.entry_archivo = tk.Entry(self.frame_archivo, width=55)
        self.entry_archivo.pack(side=tk.LEFT, padx=5)

        self.btn_examinar = tk.Button(
            self.frame_archivo,
            text="Examinar",
            command=self.seleccionar_archivo
        )
        self.btn_examinar.pack(side=tk.LEFT)

        self.label_drop = tk.Label(
            root,
            text="Seleccione un archivo o\narrastre y suelte el archivo aquí",
            width=48,
            height=5,
            relief="ridge",
            font=("Arial", 11)
        )
        self.label_drop.pack(pady=12)

        self.label_drop.drop_target_register(DND_FILES)
        self.label_drop.dnd_bind("<<Drop>>", self.archivo_soltado)

        self.lista = tk.Listbox(root, width=72, height=5, bd=2, relief="sunken")
        self.lista.pack(pady=8)

        # -----------------------------
        # PROCESAR EXCEL / CSV
        # -----------------------------
        self.btn_procesar = tk.Button(
            root,
            text="Procesar archivo",
            command=self.procesar_archivo
        )
        self.btn_procesar.pack(pady=8)

        # -----------------------------
        # CONVERSIÓN DIRECTA
        # -----------------------------
        self.lbl_convertir = tk.Label(root, text="Convertir archivo (sin procesar)")
        self.lbl_convertir.pack(pady=6)

        self.formato_convertir = tk.StringVar(value="PDF")

        self.combo_convertir = tk.OptionMenu(
            root,
            self.formato_convertir,
            "PDF",
            "WORD"
        )
        self.combo_convertir.pack(pady=4)

        self.btn_convertir = tk.Button(
            root,
            text="Convertir",
            command=self.convertir_directo
        )
        self.btn_convertir.pack(pady=6)

        # -----------------------------
        # EXPORTACIÓN
        # -----------------------------
        self.lbl_exportar = tk.Label(root, text="Exportar último archivo generado")
        self.lbl_exportar.pack(pady=6)

        self.formato_exportar = tk.StringVar(value="PDF")

        self.combo_exportar = tk.OptionMenu(
            root,
            self.formato_exportar,
            "PDF",
            "WORD",
            "CSV",
            "TXT"
        )
        self.combo_exportar.pack(pady=4)

        self.btn_exportar = tk.Button(
            root,
            text="Exportar archivo",
            command=self.exportar_generico
        )
        self.btn_exportar.pack(pady=6)

        self.aplicar_tema()

    # =============================
    # TEMA
    # =============================
    def toggle_tema(self):
        self.tema = TEMA_CLARO if self.tema == TEMA_OSCURO else TEMA_OSCURO
        self.aplicar_tema()

    
    def aplicar_tema(self):
      t = self.tema
      self.root.configure(bg=t["BG"])

      for widget in self.root.winfo_children():
 
        if isinstance(widget, tk.Frame):
            widget.configure(bg=t["BG"])

        elif isinstance(widget, tk.Label):
            widget.configure(bg=t["BG"], fg=t["FG"])

        elif isinstance(widget, tk.Button):
            widget.configure(
                bg=t["BTN_BG"],
                fg=t["BTN_FG"],
                activebackground=t["ACCENT"],
                activeforeground="white",
                relief="flat",
                bd=2
            )

        elif isinstance(widget, tk.Entry):
            widget.configure(
                bg=t["ENTRY_BG"],
                fg=t["FG"],
                insertbackground=t["FG"],
                relief="flat"
            )

        elif isinstance(widget, tk.Listbox):
            widget.configure(
                bg=t["LIST_BG"],
                fg=t["FG"],
                selectbackground=t["ACCENT"]
            )

        elif isinstance(widget, tk.OptionMenu):
            widget.configure(bg=t["BTN_BG"])
            widget["menu"].configure(
                bg=t["BTN_BG"],
                fg=t["BTN_FG"],
                activebackground=t["ACCENT"],
                activeforeground="white"
            )

      self.label_drop.configure(
        bg=t["DROP_BG"],
        fg=t["DROP_FG"]
    )


    # =============================
    # ARCHIVOS
    # =============================

    def seleccionar_archivo(self):
        ruta = filedialog.askopenfilename(title="Seleccionar archivo")
        if ruta:
            self.cargar_archivo(Path(ruta))

    def archivo_soltado(self, event):
        ruta = event.data.strip("{}")
        if os.path.isfile(ruta):
            self.cargar_archivo(Path(ruta))

    def cargar_archivo(self, archivo):
        self.archivo_seleccionado = archivo
        self.entry_archivo.delete(0, tk.END)
        self.entry_archivo.insert(0, str(archivo))
        self.lista.delete(0, tk.END)
        self.lista.insert(tk.END, archivo.name)

    # =============================
    # PROCESAR EXCEL / CSV
    # =============================

    def procesar_archivo(self):
        if not self.archivo_seleccionado:
            messagebox.showwarning("Atención", "Seleccione un archivo")
            return

        threading.Thread(target=self._procesar_thread, daemon=True).start()

    def _procesar_thread(self):
        archivo = self.archivo_seleccionado

        try:
            if archivo.suffix.lower() == ".csv":
                df = pd.read_csv(archivo)
            else:
                df = pd.read_excel(archivo)

            df = df.drop_duplicates().dropna(how="all")
            salida = archivo.with_name(f"{archivo.stem}_procesado.xlsx")
            df.to_excel(salida, index=False)

            self.ultimo_archivo = salida
            messagebox.showinfo("Listo", f"Archivo generado:\n{salida.name}")

        except Exception as e:
            messagebox.showerror("Error", str(e))

    # =============================
    # CONVERSIÓN DIRECTA
    # =============================

    def convertir_directo(self):
        if not self.archivo_seleccionado:
            messagebox.showwarning("Atención", "Seleccione un archivo")
            return

        origen = self.archivo_seleccionado
        destino = self.formato_convertir.get().lower()

        if origen.suffix == ".docx" and destino == "pdf":
            self.word_a_pdf(origen)

        elif origen.suffix == ".pdf" and destino == "word":
            self.pdf_a_word(origen)

        else:
            messagebox.showwarning("Formato no soportado", "Conversión inválida")

    def word_a_pdf(self, archivo):
        doc = Document(archivo)
        salida = archivo.with_suffix(".pdf")

        styles = getSampleStyleSheet()
        contenido = []

        for p in doc.paragraphs:
            contenido.append(Paragraph(p.text, styles["Normal"]))
            contenido.append(Spacer(1, 6))

        pdf = SimpleDocTemplate(str(salida), pagesize=A4)
        pdf.build(contenido)

        messagebox.showinfo("Listo", f"PDF generado:\n{salida.name}")

    def pdf_a_word(self, archivo):
        reader = PdfReader(str(archivo))
        salida = archivo.with_suffix(".docx")

        doc = Document()
        for page in reader.pages:
            doc.add_paragraph(page.extract_text() or "")

        doc.save(salida)
        messagebox.showinfo("Listo", f"Word generado:\n{salida.name}")

    # =============================
    # EXPORTAR
    # =============================

    def exportar_generico(self):
        if not self.ultimo_archivo:
            messagebox.showwarning("Atención", "No hay archivo procesado")
            return

        f = self.formato_exportar.get().lower()

        if f == "pdf":
            self.exportar_pdf()
        elif f == "word":
            self.exportar_word()
        elif f == "csv":
            self.exportar_csv()
        elif f == "txt":
            self.exportar_txt()

    def exportar_pdf(self):
        df = pd.read_excel(self.ultimo_archivo)
        path = self.ultimo_archivo.with_suffix(".pdf")

        doc = SimpleDocTemplate(str(path), pagesize=A4)
        styles = getSampleStyleSheet()
        data = [df.columns.tolist()] + df.values.tolist()

        tabla = Table(data)
        tabla.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ]))

        doc.build([Paragraph("Informe", styles["Title"]), tabla])
        messagebox.showinfo("Listo", f"PDF generado:\n{path.name}")

    def exportar_word(self):
        df = pd.read_excel(self.ultimo_archivo)
        path = self.ultimo_archivo.with_suffix(".docx")

        doc = Document()
        table = doc.add_table(rows=1, cols=len(df.columns))

        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = col

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

        doc.save(path)
        messagebox.showinfo("Listo", f"Word generado:\n{path.name}")

    def exportar_csv(self):
        df = pd.read_excel(self.ultimo_archivo)
        path = self.ultimo_archivo.with_suffix(".csv")
        df.to_csv(path, index=False)
        messagebox.showinfo("Listo", f"CSV generado:\n{path.name}")

    def exportar_txt(self):
        df = pd.read_excel(self.ultimo_archivo)
        path = self.ultimo_archivo.with_suffix(".txt")

        with open(path, "w", encoding="utf-8") as f:
            f.write(" | ".join(df.columns) + "\n")
            for _, row in df.iterrows():
                f.write(" | ".join(map(str, row.values)) + "\n")

        messagebox.showinfo("Listo", f"TXT generado:\n{path.name}")


# =============================
# MAIN
# =============================

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = AutomatizadorApp(root)
    root.mainloop()
